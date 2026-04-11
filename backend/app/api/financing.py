# backend/app/api/financing.py
import json
from fastapi import APIRouter, Depends, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import io

from app.core.security import get_current_user
from app.db.connection import get_conn

router = APIRouter()

# ── リクエストモデル ──────────────────────────────────────
class FiscalYear(BaseModel):
    year: str
    sales: str
    operatingProfit: str
    netAssets: str
    totalDebt: str
    cashFlow: str

class FinancingReportRequest(BaseModel):
    companyName: str
    industry: str = ""
    established: str = ""
    representative: str = ""
    capital: str = ""
    employees: str = ""
    fiscalYears: list[FiscalYear]
    currentDebt: str = ""
    monthlyRepayment: str = ""
    loanAmount: str = ""
    loanPurpose: str = ""
    repaymentPeriod: str = ""
    salesPlan1: str = ""
    salesPlan2: str = ""
    salesPlan3: str = ""
    businessDescription: str = ""

# ── 財務指標の自動計算 ────────────────────────────────────
def calc_metrics(fy: FiscalYear) -> dict:
    try:
        sales   = float(fy.sales)           or 1
        profit  = float(fy.operatingProfit) or 0
        assets  = float(fy.netAssets)       or 0
        debt    = float(fy.totalDebt)       or 1
        cf      = float(fy.cashFlow)        or 1
        return {
            "profit_margin":    round(profit / sales * 100, 1),
            "debt_repayment_y": round(debt / cf, 1),
            "equity_ratio":     round(assets / (assets + debt) * 100, 1) if assets + debt else 0,
        }
    except Exception:
        return {}

# ── AIプロンプト生成 ──────────────────────────────────────
def build_prompt(req: FinancingReportRequest) -> str:
    fy_text = "\n".join([
        f"{fy.year}：売上{fy.sales}万円 / 経常利益{fy.operatingProfit}万円 / "
        f"純資産{fy.netAssets}万円 / 借入{fy.totalDebt}万円 / 営業CF{fy.cashFlow}万円"
        for fy in req.fiscalYears if fy.year
    ])
    metrics = [calc_metrics(fy) for fy in req.fiscalYears if fy.year]

    return f"""
あなたは金融機関向け融資審査を熟知した中小企業診断士です。
以下の企業データをもとに、金融機関に提出できる水準の融資対策レポートを作成してください。

【企業概要】
会社名：{req.companyName}
業種：{req.industry} / 設立：{req.established} / 代表：{req.representative}
資本金：{req.capital}万円 / 従業員：{req.employees}名
事業内容：{req.businessDescription}

【財務データ（3期）】
{fy_text}

【財務指標（自動計算）】
{json.dumps(metrics, ensure_ascii=False, indent=2)}

【借入・融資希望】
現在借入残高：{req.currentDebt}万円 / 月次返済：{req.monthlyRepayment}万円
融資希望額：{req.loanAmount}万円 / 使途：{req.loanPurpose} / 返済期間：{req.repaymentPeriod}年

【売上計画】
1年目：{req.salesPlan1}万円 / 2年目：{req.salesPlan2}万円 / 3年目：{req.salesPlan3}万円

以下の構成でレポートを作成してください：

# 融資対策レポート

## 1. 会社概要
## 2. 財務分析（3期推移）
　- 売上・利益トレンド評価
　- 自己資本比率・流動比率・債務償還年数の評価
　- キャッシュフロー3区分分析
## 3. 財務体質の強み・課題
## 4. 事業計画
　- 市場環境と競合優位性
　- 向こう3年の売上・収支計画
　- 資金使途の詳細と必要性
## 5. 返済計画・返済シミュレーション
## 6. リスクと対応策
## 7. 総合評価コメント（金融機関への訴求ポイント）
"""

# ── PDF生成 ───────────────────────────────────────────────
def generate_pdf(content: str, company_name: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    pdfmetrics.registerFont(UnicodeCIDFont("HeiseiMin-W3"))
    pdfmetrics.registerFont(UnicodeCIDFont("HeiseiKakuGo-W5"))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=20*mm, bottomMargin=20*mm,
    )

    styles = getSampleStyleSheet()
    style_h1 = ParagraphStyle("h1", fontName="HeiseiKakuGo-W5", fontSize=16,
                               spaceAfter=8, textColor=colors.HexColor("#1e3a5f"))
    style_h2 = ParagraphStyle("h2", fontName="HeiseiKakuGo-W5", fontSize=12,
                               spaceAfter=6, spaceBefore=12,
                               textColor=colors.HexColor("#2563eb"))
    style_body = ParagraphStyle("body", fontName="HeiseiMin-W3", fontSize=10,
                                leading=16, spaceAfter=4)

    story = []
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 4))
        elif line.startswith("# "):
            story.append(Paragraph(line[2:], style_h1))
            story.append(HRFlowable(width="100%", thickness=1,
                                    color=colors.HexColor("#2563eb")))
            story.append(Spacer(1, 4))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:], style_h2))
        else:
            story.append(Paragraph(line.replace("　", "&nbsp;&nbsp;"), style_body))

    doc.build(story)
    return buf.getvalue()

# ── Word生成 ──────────────────────────────────────────────
def generate_docx(content: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("# "):
            h = doc.add_heading(line[2:], level=1)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
        elif line.startswith("## "):
            h = doc.add_heading(line[3:], level=2)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
        else:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(10.5)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ── メインエンドポイント ──────────────────────────────────
@router.post("/generate")
async def generate_report(
    req: FinancingReportRequest,
    format: str = Query("pdf", pattern="^(pdf|docx)$"),
    user: dict = Depends(get_current_user),
):
    import anthropic

    client = anthropic.AsyncAnthropic()
    prompt = build_prompt(req)

    message = await client.messages.create(
        model="claude-opus-4-5",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )
    report_text = message.content[0].text

    if format == "pdf":
        data     = generate_pdf(report_text, req.companyName)
        media    = "application/pdf"
        filename = f"融資対策レポート_{req.companyName}.pdf"
    else:
        data     = generate_docx(report_text)
        media    = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"融資対策レポート_{req.companyName}.docx"

    return StreamingResponse(
        io.BytesIO(data),
        media_type=media,
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"},
    )
